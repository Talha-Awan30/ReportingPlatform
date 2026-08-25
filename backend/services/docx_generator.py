"""Word report generation.

Two entry points:

  * `generate(report, spec)`     - one unit's report on its own.
  * `generate_set(record, spec)` - a whole visit: the shared title page once,
                                   then every unit's report in sequence.

Both fill the module's approved .docx template with docxtpl when one is present.
If a module has not supplied a template yet, a readable fallback document is
built with python-docx that follows the same layout as the approved format, so
the workflow works end to end before the template arrives.
"""
import logging
import os
from datetime import datetime

from models import PhotoKind
from services.storage import generated_dir, photo_path, set_photo_path

log = logging.getLogger(__name__)


def template_path(spec):
    """Absolute path to a module's Word template, or None if it is not there."""
    if not spec or not spec.root_path:
        return None
    path = os.path.join(spec.root_path, "templates", spec.docx_template)
    return path if os.path.exists(path) else None


# ============================================================================
# CONTEXT
# ============================================================================
def build_context(report, spec, options=None):
    """The flat dictionary one unit's Word template renders against."""
    job = report.job
    equipment = report.equipment
    client = job.client if job else None
    options = options or {}
    answers = report.data or {}

    def label_for(field, raw_value):
        """Turn a stored option value into its approved report wording."""
        if field.kind != "dropdown" or raw_value in (None, ""):
            return raw_value
        bucket = options.get(field.options_key) or {}
        for option in bucket.get("options", []):
            if option["value"] == raw_value:
                return option.get("report_text") or option["label"]
        return raw_value

    def render_field(field):
        entry = answers.get(field.key) or {}
        if not isinstance(entry, dict):
            entry = {"value": entry}
        raw = entry.get("value")
        return {
            "key": field.key,
            "label": field.label,
            "value": raw,
            "display": label_for(field, raw) if field.kind == "dropdown" else raw,
            "remarks": entry.get("remarks", ""),
        }

    # --- the check-list, grouped by section ---------------------------------
    sections = []
    for section in spec.sections if spec else []:
        rows = []
        for checkpoint in section.checkpoints:
            row = render_field(checkpoint)
            row["result"] = row["display"]
            rows.append(row)
        sections.append({"key": section.key, "title": section.title, "rows": rows})

    # --- the particulars table and the closing findings ---------------------
    unit_details = [render_field(f) for f in (spec.unit_details if spec else [])]
    conclusion_rows = [render_field(f) for f in (spec.conclusion if spec else [])]

    # --- photographs, bucketed into the module's named slots ----------------
    by_slot = {}
    for photo in report.photos_of(PhotoKind.INSPECTION):
        by_slot.setdefault(photo.checkpoint_key or "", []).append(_photo(photo))
    photo_slots = [
        {"key": slot.key, "label": slot.label, "photos": by_slot.get(slot.key, [])}
        for slot in (spec.photo_slots if spec else [])
    ]

    return {
        "report_number": report.report_number,
        "sequence": report.sequence,
        "revision": report.revision,
        "status": report.status.label,
        "module_name": spec.name if spec else report.module_slug,
        "inspection_date": _fmt(report.inspection_date),
        "next_inspection_date": _fmt(report.next_inspection_date),
        "certificate_expiry_date": _fmt(report.certificate_expiry_date),
        "overall_result": (report.overall_result or "").title(),
        "comments": report.comments or "",
        "job_number": job.job_number if job else "",
        "site_name": job.site_name if job else "",
        "site_address": job.site_address if job else "",
        "client_name": client.name if client else "",
        "client_address": client.address if client else "",
        "contact_name": job.contact.name if job and job.contact else "",
        "equipment_tag": equipment.tag_number if equipment else "",
        "equipment_type": equipment.equipment_type.name if equipment and equipment.equipment_type else "",
        "serial_number": equipment.serial_number if equipment else "",
        "manufacturer": equipment.manufacturer if equipment else "",
        "model": equipment.model if equipment else "",
        "capacity": equipment.capacity if equipment else "",
        "swl": equipment.swl if equipment else "",
        "location": equipment.location if equipment else "",
        "inspector_name": report.inspector.full_name if report.inspector else "",
        "reviewer_name": report.reviewer.full_name if report.reviewer else "",
        "generated_on": datetime.now().strftime("%d %B %Y"),
        "unit_details": unit_details,
        "conclusion": conclusion_rows,
        "photo_slots": photo_slots,
        "sections": sections,
        "front_page_photos": [_photo(p) for p in report.photos_of(PhotoKind.FRONT_PAGE)],
        "inspection_photos": [_photo(p) for p in report.photos_of(PhotoKind.INSPECTION)],
    }


def _set_context(record, spec, options):
    """Title-page values plus one context per unit."""
    return {
        "set_number": record.set_number,
        "unit_count": len(record.reports),
        "unit_noun": spec.unit_noun if spec else "unit",
        "unit_noun_plural": spec.unit_noun_plural if spec else "units",
        "title_page": _flatten(record.title_page or {}),
        "title_page_photos": [
            {
                "slot": p.slot_key,
                "path": set_photo_path(p),
                "caption": p.caption or p.original_name or "",
            }
            for p in record.photos
        ],
        "units": [build_context(r, spec, options) for r in record.reports],
        "generated_on": datetime.now().strftime("%d %B %Y"),
    }


def _flatten(stored):
    """Title-page values are stored as {key: {value: ...}} - unwrap them."""
    out = {}
    for key, entry in (stored or {}).items():
        out[key] = entry.get("value") if isinstance(entry, dict) else entry
    return out


# ============================================================================
# ENTRY POINTS
# ============================================================================
def generate(report, spec, options=None):
    """Render one unit's report to .docx; returns its name inside GENERATED_FOLDER."""
    context = build_context(report, spec, options)
    out_name = f"{report.report_number or f'report-{report.id}'}-rev{report.revision}.docx"
    out_path = os.path.join(generated_dir(), out_name)

    source = template_path(spec)
    if source:
        _render_template(source, out_path, context)
    else:
        log.info("No Word template for module '%s' - building fallback document", report.module_slug)
        _render_fallback(out_path, context, spec)

    return out_name


def generate_set(record, spec):
    """Render a whole visit: the title page once, then every unit in sequence."""
    from modules.blueprint_factory import resolve_options

    options = resolve_options(spec) if spec else {}
    out_name = f"{record.set_number or f'set-{record.id}'}.docx"
    out_path = os.path.join(generated_dir(), out_name)

    source = template_path(spec)
    if source:
        _render_set_template(source, out_path, record, spec, options)
    else:
        log.info(
            "No Word template for module '%s' - building fallback set document", record.module_slug
        )
        _render_set_fallback(out_path, record, spec, options)

    return out_name


# ============================================================================
# TEMPLATE RENDERING (docxtpl)
# ============================================================================
def _render_template(source, out_path, context):
    from docx.shared import Mm
    from docxtpl import DocxTemplate, InlineImage

    doc = DocxTemplate(source)
    _bind_images(doc, context, InlineImage, Mm)
    doc.render(context)
    doc.save(out_path)


def _render_set_template(source, out_path, record, spec, options):
    from docx.shared import Mm
    from docxtpl import DocxTemplate, InlineImage

    doc = DocxTemplate(source)
    context = _set_context(record, spec, options)

    for item in context["title_page_photos"]:
        if item["path"] and os.path.exists(item["path"]):
            item["image"] = InlineImage(doc, item["path"], width=Mm(120))
    for unit in context["units"]:
        _bind_images(doc, unit, InlineImage, Mm)

    doc.render(context)
    doc.save(out_path)


def _bind_images(doc, context, InlineImage, Mm):
    """Photos have to become InlineImage objects bound to this specific document."""
    for key in ("front_page_photos", "inspection_photos"):
        for item in context.get(key, []):
            if item["path"] and os.path.exists(item["path"]):
                item["image"] = InlineImage(doc, item["path"], width=Mm(80))
    for slot in context.get("photo_slots", []):
        for item in slot["photos"]:
            if item["path"] and os.path.exists(item["path"]):
                item["image"] = InlineImage(doc, item["path"], width=Mm(75))


# ============================================================================
# FALLBACK RENDERING (python-docx), laid out like the approved format
# ============================================================================
def _render_fallback(out_path, context, spec=None):
    """A single unit's report, used until the approved template arrives."""
    from docx import Document
    from docx.shared import RGBColor

    doc = Document()
    sgs_blue = RGBColor(0x3C, 0x51, 0x5B)
    _write_unit(doc, context, spec, 1, 1, sgs_blue)
    doc.save(out_path)


def _render_set_fallback(out_path, record, spec, options):
    """Title page + one full report per unit."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm, Pt, RGBColor

    doc = Document()
    context = _set_context(record, spec, options)
    tp = context["title_page"]
    sgs_blue = RGBColor(0x3C, 0x51, 0x5B)

    # ------------------------------------------------------------ TITLE PAGE
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(tp.get("report_title") or "Third Party Inspection Report")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = sgs_blue

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("Inspection Report")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = sgs_blue

    for item in context["title_page_photos"]:
        if item["path"] and os.path.exists(item["path"]):
            doc.add_picture(item["path"], width=Mm(140))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    table = doc.add_table(rows=0, cols=3)
    table.style = "Table Grid"
    for label, key in [
        ("Client", "client"),
        ("Client Ref", "client_ref"),
        ("Client Contact Person", "client_contact_person"),
        ("SGS Ref. No.", "sgs_ref_no"),
        ("Inspected by", "inspected_by"),
        ("Site", "site"),
        ("Equipment Identification", "equipment_identification"),
        ("Survey Date", "survey_date"),
        ("QR Code", "qr_code"),
    ]:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = ":"
        cells[2].text = str(tp.get(key) or "")

    doc.add_paragraph()
    signoff = doc.add_table(rows=2, cols=2)
    signoff.style = "Table Grid"
    signoff.rows[0].cells[0].text = "Prepared by :"
    signoff.rows[0].cells[1].text = "Reviewed by:"
    signoff.rows[1].cells[0].text = _join(tp.get("prepared_by_name"), tp.get("prepared_by_designation"))
    signoff.rows[1].cells[1].text = _join(tp.get("reviewed_by_name"), tp.get("reviewed_by_designation"))

    # --------------------------------------------------------- ONE PER UNIT
    total = len(context["units"])
    for index, unit in enumerate(context["units"], start=1):
        doc.add_page_break()
        _write_unit(doc, unit, spec, index, total, sgs_blue)

    doc.save(out_path)


def _write_unit(doc, unit, spec, index, total, sgs_blue):
    """One unit's full report: particulars, check-list, findings, photographs."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm, Pt

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Examination Report")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = sgs_blue

    noun = (spec.unit_noun if spec else "unit").title()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = f"{noun} {index} of {total}  -  {unit['report_number']}" if total > 1 else unit["report_number"]
    sub.add_run(f"Of  {caption}").font.size = Pt(12)

    # --- particulars table --------------------------------------------------
    doc.add_paragraph()
    details = unit.get("unit_details") or []
    if details:
        table = doc.add_table(rows=0, cols=3)
        table.style = "Table Grid"
        for row in details:
            cells = table.add_row().cells
            cells[0].text = row["label"]
            cells[1].text = ":"
            cells[2].text = str(row["display"] or "")

    doc.add_paragraph()
    doc.add_paragraph("INSPECTION FINDINGS").runs[0].bold = True
    doc.add_paragraph(
        "Following points were checked during the inspection activity to ensure safe "
        "operational condition along with proper working of safety devices."
    )

    # --- check-list, two columns exactly as the approved format -------------
    for section in unit["sections"]:
        doc.add_paragraph()
        doc.add_paragraph(section["title"].upper()).runs[0].bold = True
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        header = table.rows[0].cells
        header[0].text = "Check Points"
        header[1].text = "Remarks/Recommendations"
        for cell in header:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for entry in section["rows"]:
            cells = table.add_row().cells
            cells[0].text = entry["label"]
            # The approved clause for the chosen option, plus any typed remark.
            parts = [p for p in (entry.get("result"), entry.get("remarks")) if p]
            cells[1].text = "  ".join(str(p) for p in parts)

    # --- findings & conclusion ---------------------------------------------
    for row in unit.get("conclusion") or []:
        doc.add_paragraph()
        doc.add_paragraph(f"{row['label'].upper()}:").runs[0].bold = True
        doc.add_paragraph(str(row["display"] or ""))

    if unit.get("comments"):
        doc.add_paragraph()
        doc.add_paragraph("OBSERVATIONS & RECOMMENDATIONS:").runs[0].bold = True
        doc.add_paragraph(unit["comments"])

    # --- photographic presentation -----------------------------------------
    slots = unit.get("photo_slots") or []
    if any(slot["photos"] for slot in slots):
        doc.add_page_break()
        doc.add_paragraph("Photographic Presentation").runs[0].bold = True
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"

        for start in range(0, len(slots), 2):
            pair = slots[start:start + 2]
            label_row = table.add_row().cells
            image_row = table.add_row().cells
            for column, slot in enumerate(pair):
                label_row[column].text = slot["label"]
                cell = image_row[column]
                path = slot["photos"][0]["path"] if slot["photos"] else None
                if path and os.path.exists(path):
                    cell.paragraphs[0].add_run().add_picture(path, width=Mm(75))
                else:
                    cell.text = "-"
    elif unit.get("inspection_photos"):
        doc.add_page_break()
        doc.add_paragraph("Photographic Presentation").runs[0].bold = True
        for item in unit["inspection_photos"]:
            if item["path"] and os.path.exists(item["path"]):
                doc.add_picture(item["path"], width=Mm(120))
                if item["caption"]:
                    doc.add_paragraph(item["caption"]).runs[0].font.size = Pt(9)


# ============================================================================
# HELPERS
# ============================================================================
def _photo(photo):
    return {
        "path": photo_path(photo),
        "caption": photo.caption or photo.original_name or "",
        "checkpoint_key": photo.checkpoint_key,
    }


def _fmt(value):
    return value.strftime("%d %b %Y") if value else ""


def _join(*parts):
    return "\n".join(str(p) for p in parts if p)
