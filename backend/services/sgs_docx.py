"""Fill the approved SGS Word document in place.

Rather than rebuilding the report layout in code, this opens the client's own
`Elevator Report Format.docx` and writes the inspector's values into the cells
and paragraphs that are already there. The output is therefore the approved
format - the same fonts, borders, headers, footers and SGS letterhead - with
the blanks filled in.

The template holds BOTH halves of the document, separated by a Word section
break:

    [ title page ]  t0 header · t1 client details · t2 prepared/reviewed
    -- section break --
    [ report body ] signature · particulars · 9 check-lists · load test ·
                    conclusion · photographic presentation

`split_template()` cuts it at that break, so one visit produces:

    title page (once)  +  report body (repeated per lift)

Tables are located by their own header text, not by position, so the template
can be re-ordered or extended without breaking the mapping.
"""
import logging
import os

from docx.oxml.ns import qn

log = logging.getLogger(__name__)

# Text that identifies where the report body begins.
BODY_MARKER = "examination report"
TITLE_MARKER = "third party inspection"


# ---------------------------------------------------------------- primitives
def set_text(paragraph, text):
    """Replace a paragraph's text, keeping the first run's formatting."""
    text = "" if text is None else str(text)
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell(cell, text, like=None):
    """Write into a table cell, keeping its existing paragraph formatting.

    The answer cells in the approved document are empty, so there is no run to
    inherit from. Pass `like` - normally the check-point cell in the same row -
    and the new text picks up that cell's font, size and colour, so the answer
    sits in the box looking exactly like the text beside it.
    """
    text = "" if text is None else str(text)
    if not cell.paragraphs:
        cell.add_paragraph(text)
        return

    had_runs = bool(cell.paragraphs[0].runs)
    set_text(cell.paragraphs[0], text)

    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)

    if not had_runs and like is not None:
        _copy_run_format(like, cell.paragraphs[0])


def _copy_run_format(source_cell, target_paragraph):
    """Copy font name/size/colour from a source cell's first run."""
    source_run = next(
        (r for p in source_cell.paragraphs for r in p.runs if r.text.strip()), None
    )
    if source_run is None or not target_paragraph.runs:
        return

    target = target_paragraph.runs[0]
    try:
        if source_run.font.name:
            target.font.name = source_run.font.name
        if source_run.font.size:
            target.font.size = source_run.font.size
        if source_run.font.color and source_run.font.color.rgb:
            target.font.color.rgb = source_run.font.color.rgb
        # Answers are plain text - never inherit a heading's bold/italic.
        target.bold = False
        target.italic = False
    except Exception:  # noqa: BLE001 - formatting is best-effort
        log.debug("Could not copy run formatting", exc_info=True)


def replace_in_paragraph(paragraph, mapping):
    """Substitute placeholder fragments while keeping the run's formatting."""
    full = "".join(run.text for run in paragraph.runs)
    if not full:
        return
    replaced = full
    for needle, value in mapping.items():
        if needle in replaced:
            replaced = replaced.replace(needle, "" if value is None else str(value))
    if replaced != full:
        set_text(paragraph, replaced)


def add_image_to_cell(cell, path, width_mm=75):
    """Append a picture underneath whatever label the cell already carries."""
    from docx.shared import Mm

    if not path or not os.path.exists(path):
        return False
    paragraph = cell.add_paragraph()
    paragraph.add_run().add_picture(path, width=Mm(width_mm))
    return True


# ------------------------------------------------------------- table finding
def _cell_text(cell):
    return " ".join(cell.text.split()).strip()


def _row_texts(table, row=0):
    if not table.rows:
        return []
    return [_cell_text(c) for c in table.rows[row].cells]


def find_tables(doc):
    """Locate every block of the report body by its own header text.

    Returns a dict of the tables the filler needs. Anything it cannot find is
    simply absent, and that part of the report is skipped.
    """
    found = {"checklists": []}

    for table in doc.tables:
        header = _row_texts(table)
        first = header[0] if header else ""
        lowered = first.lower()

        if len(header) == 2 and header[1].lower().startswith("remarks"):
            found["checklists"].append(table)
        elif len(header) == 2 and lowered == "description" and header[1].lower() == "result":
            found["load_test"] = table
        elif len(header) == 2 and lowered == "particulars":
            found["signature"] = table
        elif len(header) == 3 and lowered.startswith("name of client"):
            found["particulars"] = table
        elif len(header) == 2 and lowered == "cabin":
            found["photos"] = table
        elif len(header) == 3 and lowered.startswith("next examination"):
            found["conclusion"] = table
        elif any(
            _cell_text(c).lower().startswith("next examination")
            for r in table.rows
            for c in r.cells
        ) and len(header) == 3:
            found["conclusion"] = table

    return found


# ------------------------------------------------------- splitting the parts
def _split_index(doc):
    """Index of the first body child that belongs to the report body.

    The cut is the paragraph carrying the section break that ends the title
    page. Falls back to the "Examination Report" heading if there is no break.
    """
    from docx.text.paragraph import Paragraph

    body = doc.element.body
    children = list(body)

    marker_index = None
    for index, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue
        if Paragraph(child, doc).text.strip().lower().startswith(BODY_MARKER):
            marker_index = index
            break

    if marker_index is None:
        return None

    # The last section break before the heading ends the title page.
    for index in range(marker_index - 1, -1, -1):
        child = children[index]
        if child.tag == qn("w:p") and child.find(".//" + qn("w:sectPr")) is not None:
            return index + 1

    return marker_index


def has_embedded_title_page(path):
    """True when this template carries its own cover page."""
    from docx import Document

    try:
        doc = Document(path)
    except Exception:
        log.exception("Could not open template %s", path)
        return False

    if _split_index(doc) is None:
        return False
    return any(TITLE_MARKER in _cell_text(c).lower() for t in doc.tables for r in t.rows for c in r.cells)


def _drop(children):
    for child in children:
        parent = child.getparent()
        if parent is not None:
            parent.remove(child)


def load_title_part(path):
    """The template with the report body removed - the cover page alone."""
    from docx import Document

    doc = Document(path)
    split = _split_index(doc)
    if split is None:
        return doc

    body = doc.element.body
    children = list(body)

    # Everything from the split to the body's own sectPr belongs to the report.
    tail = [c for c in children[split:] if c.tag != qn("w:sectPr")]
    _drop(tail)

    # The title page's section properties live on its last paragraph. Promote
    # them to the body so the cover keeps its own page setup, headers and
    # footers, and no empty trailing section is left behind.
    remaining = list(body)
    for child in reversed(remaining):
        if child.tag != qn("w:p"):
            continue
        sect = child.find(".//" + qn("w:sectPr"))
        if sect is None:
            continue
        body_sect = body.find(qn("w:sectPr"))
        if body_sect is not None:
            body.remove(body_sect)
        sect.getparent().remove(sect)
        body.append(sect)
        break

    return doc


def load_unit_part(path):
    """The template with the cover page removed - one lift's report body."""
    from docx import Document

    doc = Document(path)
    split = _split_index(doc)
    if split is None:
        return doc

    _drop(list(doc.element.body)[:split])
    return doc


# --------------------------------------------------------------- title page
TITLE_FIELDS = {
    "client": "client",
    "client ref": "client_ref",
    "client contact person": "client_contact_person",
    "sgs ref. no.": "sgs_ref_no",
    "sgs ref no.": "sgs_ref_no",
    "sgs ref no": "sgs_ref_no",
    "inspected by": "inspected_by",
    "site": "site",
    "equipment identification": "equipment_identification",
    "survey date": "survey_date",
    "qr code": "qr_code",
}


def fill_title_page(doc_or_path, values, photos=None, spec=None):
    """Write the visit's details onto the approved cover page."""
    doc = _as_doc(doc_or_path)

    # The template's own picture frames, captured NOW. A picture written into a
    # table cell later (the QR code) would otherwise join `inline_shapes` and,
    # because that table sits above the frames, shift every later match by one.
    frames = list(doc.inline_shapes)

    # A photo slot may claim a row of the details table (the QR Code), in
    # which case that row takes an image instead of text.
    slot_by_label = {
        slot.cell_label.strip().lower(): slot
        for slot in (spec.title_page_photos if spec else [])
        if slot.cell_label
    }
    photo_by_slot = {p.get("slot"): p for p in (photos or []) if p.get("slot")}

    # The cover heading, carried in the first table's left cell.
    if values.get("report_title"):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if TITLE_MARKER in _cell_text(cell).lower():
                        set_cell(cell, values["report_title"])

    # The details table: label | : | value.
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 3:
                continue
            label = _cell_text(row.cells[0]).lower().rstrip(":")

            if label in slot_by_label:
                # This row holds a picture, not text.
                photo = photo_by_slot.get(slot_by_label[label].key)
                if photo and photo.get("path") and os.path.exists(photo["path"]):
                    _set_cell_image(row.cells[2], photo["path"])
                continue

            key = TITLE_FIELDS.get(label)
            if key:
                set_cell(row.cells[2], values.get(key, ""), like=row.cells[0])

    # Prepared by / Reviewed by.
    for table in doc.tables:
        for index, row in enumerate(table.rows):
            if len(row.cells) != 2:
                continue
            if "prepared by" in _cell_text(row.cells[0]).lower() and index + 1 < len(table.rows):
                target = table.rows[index + 1].cells
                set_cell(target[0], _join(values.get("prepared_by_name"),
                                          values.get("prepared_by_designation")))
                set_cell(target[1], _join(values.get("reviewed_by_name"),
                                          values.get("reviewed_by_designation")))

    frame_slots = [s for s in (spec.title_page_photos if spec else []) if not s.cell_label]
    apply_cover_photos(doc, frames, frame_slots, photo_by_slot, photos)
    _move_frames_below_header(doc)
    return doc


def _move_frames_below_header(doc):
    """Put the cover photographs in the gap under the report header.

    The template leaves them below the client details table, which prints them
    low on the page with a blank band above. Moving the paragraph that carries
    them directly after the header table fills that space instead.
    """
    body = doc.element.body

    header = next(
        (
            child
            for child in body
            if child.tag == qn("w:tbl")
            and TITLE_MARKER in " ".join(child.itertext()).lower()
        ),
        None,
    )
    frame_para = next(
        (
            child
            for child in body
            if child.tag == qn("w:p") and child.find(".//" + qn("a:blip")) is not None
        ),
        None,
    )

    if header is None or frame_para is None:
        return False

    # lxml moves the element rather than copying it.
    header.addnext(frame_para)

    # The template left blank paragraphs where the pictures used to be. With
    # the frames moved up they would print as a band of white, so drop the
    # empty ones between the photographs and the next table.
    node = frame_para.getnext()
    while node is not None and node.tag == qn("w:p"):
        following = node.getnext()
        empty = (
            not "".join(node.itertext()).strip()
            and node.find(".//" + qn("a:blip")) is None
            and node.find(".//" + qn("w:sectPr")) is None
        )
        if not empty:
            break
        node.getparent().remove(node)
        node = following

    return True


def apply_cover_photos(doc, frames, frame_slots, photo_by_slot, all_photos=None):
    """Fill the cover's picture frames, and clear the ones left unused.

    Each frame belongs to one manifest slot, in declared order, so the
    "Cover photograph" upload always lands in the first frame. A frame whose
    slot was not filled in is removed outright, so the template's sample
    photograph never survives into a real report.
    """
    if frame_slots:
        pairs = [(frame, photo_by_slot.get(slot.key)) for frame, slot in zip(frames, frame_slots)]
    else:
        # No manifest to go on: fall back to filling the frames in order.
        usable = [p for p in (all_photos or []) if p.get("path") and os.path.exists(p["path"])]
        pairs = list(zip(frames, usable + [None] * len(frames)))

    for frame, photo in pairs:
        if photo and photo.get("path") and os.path.exists(photo["path"]):
            _swap_frame_image(doc, frame, photo["path"])
        else:
            _remove_frame(frame)


def _set_cell_image(cell, path, width_mm=18):
    """Replace a cell's contents with a single picture, e.g. the QR code."""
    from docx.shared import Mm

    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    paragraph.add_run().add_picture(path, width=Mm(width_mm))


# The cover photographs are sized to fit inside this box, keeping their own
# proportions. Two of them then sit comfortably side by side on the page.
COVER_MAX_W_MM = 55
COVER_MAX_H_MM = 42


def _swap_frame_image(doc, shape, path):
    """Point one picture frame at a new image, upright and sized to fit."""
    from docx.shared import Mm

    try:
        # A fresh image part carries the correct content type and extension,
        # which swapping the raw bytes of the template's part would not.
        rid, image = doc.part.get_or_add_image(path)
        pic = shape._inline.graphic.graphicData.pic
        pic.blipFill.blip.embed = rid

        # The template's frames are rotated 90 degrees to suit the sample
        # photographs. An uploaded picture is already the right way up, so the
        # rotation and any flip have to go or it prints on its side.
        _clear_rotation(pic)

        # Fit inside the box rather than filling the template's frame, so a
        # portrait photo is not blown up to half a page.
        ratio = (image.height / image.width) if image.width else 1
        width = Mm(COVER_MAX_W_MM)
        height = int(width * ratio)
        if height > Mm(COVER_MAX_H_MM):
            height = Mm(COVER_MAX_H_MM)
            width = int(height / ratio) if ratio else Mm(COVER_MAX_W_MM)

        shape.width = int(width)
        shape.height = int(height)
        return True
    except Exception:  # noqa: BLE001 - a bad image must not lose the report
        log.exception("Could not place cover photograph %s", path)
        return False


def _clear_rotation(pic):
    """Drop any rotation or flip the template applied to a picture frame."""
    for xfrm in pic.iter(qn("a:xfrm")):
        for attribute in ("rot", "flipH", "flipV"):
            if xfrm.get(attribute) is not None:
                del xfrm.attrib[attribute]


def _remove_frame(shape):
    """Delete an unused picture frame so no sample image is left behind."""
    try:
        element = shape._inline
        # Walk up to the run that carries the drawing and drop the whole run.
        node = element
        while node is not None and node.tag != qn("w:r"):
            node = node.getparent()
        target = node if node is not None else element
        parent = target.getparent()
        if parent is not None:
            parent.remove(target)
        return True
    except Exception:  # noqa: BLE001
        log.exception("Could not remove an unused cover frame")
        return False


def _set_cell_image(cell, path, width_mm=18):
    """Replace a cell's contents with a single picture, e.g. the QR code."""
    from docx.shared import Mm

    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    paragraph.add_run().add_picture(path, width=Mm(width_mm))


def replace_cover_photos(doc, photos):
    """Swap the template's cover photographs for the inspector's own.

    The template already carries the picture frames in the right place at the
    top of the cover page, so the uploaded images replace those in position
    rather than being appended to the end of the document. The frame's width is
    kept and the height recomputed from the new image, so nothing is stretched
    and the side-by-side layout survives.
    """
    usable = [p for p in (photos or []) if p.get("path") and os.path.exists(p["path"])]
    if not usable:
        return 0

    shapes = list(doc.inline_shapes)
    replaced = 0

    for shape, photo in zip(shapes, usable):
        try:
            # A fresh image part carries the correct content type and extension,
            # which swapping the raw bytes of the template's part would not.
            rid, image = doc.part.get_or_add_image(photo["path"])
            shape._inline.graphic.graphicData.pic.blipFill.blip.embed = rid

            # Keep the template's frame width; take the height from the new
            # image so it is placed, not stretched.
            if image.width:
                shape.height = int(shape.width * (image.height / image.width))

            replaced += 1
        except Exception:  # noqa: BLE001 - a bad image must not lose the report
            log.exception("Could not replace cover photograph %s", photo.get("path"))

    if len(usable) > len(shapes):
        log.info(
            "%s cover photo(s) uploaded but the template has only %s frame(s) - extras ignored",
            len(usable),
            len(shapes),
        )
    return replaced


# -------------------------------------------------------------- unit report
def fill_unit_report(doc_or_path, unit, spec, title_page=None):
    """Write one lift's answers onto the approved Examination Report."""
    doc = _as_doc(doc_or_path)
    title_page = title_page or {}
    details = {row["key"]: row for row in unit.get("unit_details") or []}

    def detail(key):
        row = details.get(key) or {}
        return row.get("display") or row.get("value") or ""

    identification = detail("identification") or unit.get("equipment_tag") or ""
    client_name = detail("name_of_client") or unit.get("client_name") or title_page.get("client", "")
    site = title_page.get("site") or unit.get("site_name") or ""
    inspection_date = detail("inspection_date") or unit.get("inspection_date") or ""
    status = detail("status")

    # --- the cover paragraphs of the Examination Report ---------------------
    for paragraph in doc.paragraphs:
        text = paragraph.text
        stripped = text.strip()
        if "(Equipment I.D)" in text:
            replace_in_paragraph(paragraph, {"(Equipment I.D)": identification})
        elif stripped.startswith("Inspection Date:"):
            set_text(paragraph, f"Inspection Date: {inspection_date}")
        elif stripped.startswith("Status:"):
            set_text(paragraph, f"Status: {status}")
        elif stripped.startswith("SGS Order No:"):
            set_text(paragraph, f"SGS Order No: {title_page.get('sgs_ref_no', '')}")
        elif stripped.startswith("Client:"):
            set_text(paragraph, f"Client: {client_name}")
        elif "(Client Name)" in text or "(Site Location)" in text:
            replace_in_paragraph(paragraph, {"(Client Name)": client_name, "(Site Location)": site})

    tables = find_tables(doc)

    # --- signature block ----------------------------------------------------
    if "signature" in tables:
        for row in tables["signature"].rows:
            if len(row.cells) < 2:
                continue
            label = _cell_text(row.cells[0]).lower().rstrip(":")
            if label == "name":
                set_cell(row.cells[1], title_page.get("inspected_by") or unit.get("inspector_name", ""),
                         like=row.cells[0])
            elif label == "date":
                set_cell(row.cells[1], inspection_date, like=row.cells[0])

    # --- particulars table, positional against spec.unit_details -----------
    if "particulars" in tables:
        for row, field in zip(tables["particulars"].rows, spec.unit_details):
            if len(row.cells) >= 3:
                value = details.get(field.key, {})
                set_cell(row.cells[2], value.get("display") or value.get("value") or "",
                         like=row.cells[0])

    # --- check-point tables, one per manifest section ----------------------
    unit_sections = {s["key"]: s for s in unit.get("sections") or []}
    checklist_sections = [s for s in spec.sections if s.key in unit_sections]

    for table, section in zip(tables["checklists"], checklist_sections):
        answers = {row["key"]: row for row in (unit_sections.get(section.key) or {}).get("rows") or []}
        for row, checkpoint in zip(table.rows[1:], section.checkpoints):
            if len(row.cells) < 2:
                continue
            set_cell(row.cells[1], _remark(answers.get(checkpoint.key, {})), like=row.cells[0])

    # --- load test exercise -------------------------------------------------
    load_section = next((s for s in spec.sections if s.key == "load_test"), None)
    if "load_test" in tables and load_section:
        rows = {r["key"]: r for r in (unit_sections.get(load_section.key) or {}).get("rows") or []}
        dropdowns = [cp for cp in load_section.checkpoints if cp.kind == "dropdown"]
        for row, checkpoint in zip(tables["load_test"].rows[1:], dropdowns):
            if len(row.cells) >= 2:
                set_cell(row.cells[1], _remark(rows.get(checkpoint.key, {})), like=row.cells[0])

    # --- next examination + conclusion -------------------------------------
    conclusion = {row["key"]: row for row in unit.get("conclusion") or []}
    if "conclusion" in tables:
        for row in tables["conclusion"].rows:
            if len(row.cells) < 3:
                continue
            label = _cell_text(row.cells[0]).lower()
            if label.startswith("next examination"):
                entry = conclusion.get("next_examination_due", {})
                set_cell(row.cells[2], entry.get("display") or entry.get("value") or "",
                         like=row.cells[0])
            elif label.startswith("conclusion"):
                entry = conclusion.get("conclusion", {})
                set_cell(row.cells[2], entry.get("display") or entry.get("value") or "",
                         like=row.cells[0])

    # --- major / minor findings, written under their headings ---------------
    for paragraph in doc.paragraphs:
        upper = paragraph.text.strip().upper()
        if upper.startswith("MAJOR FINDINGS"):
            entry = conclusion.get("major_findings", {})
            set_text(paragraph, f"MAJOR FINDINGS: {entry.get('value') or ''}")
        elif upper.startswith("MINOR FINDINGS"):
            entry = conclusion.get("minor_findings", {})
            set_text(paragraph, f"MINOR FINDINGS (AREA OF IMPROVEMENT): {entry.get('value') or ''}")

    # --- photographic presentation, into the existing grid -----------------
    if "photos" in tables:
        by_slot = {slot["key"]: slot for slot in unit.get("photo_slots") or []}
        slots = list(spec.photo_slots)
        position = 0
        for row in tables["photos"].rows:
            for cell in row.cells:
                if position >= len(slots):
                    break
                photos = (by_slot.get(slots[position].key) or {}).get("photos") or []
                if photos:
                    add_image_to_cell(cell, photos[0]["path"])
                position += 1

    return doc


# ------------------------------------------------------------------ compose
def compose(parts, out_path):
    """Merge the filled documents into one file, preserving every style."""
    from docxcompose.composer import Composer

    if not parts:
        raise ValueError("Nothing to compose")

    composer = Composer(parts[0])
    for part in parts[1:]:
        composer.append(part)
    composer.save(out_path)
    return out_path


# ------------------------------------------------------------------ helpers
def _as_doc(doc_or_path):
    from docx import Document

    return Document(doc_or_path) if isinstance(doc_or_path, str) else doc_or_path


def _remark(answer):
    """The cell text for one check point.

    Just the option's short label - "Satisfactory" - so the box stays a single
    line, followed by the inspector's own remark only when they typed one.
    """
    if not answer:
        return ""
    word = answer.get("short") or answer.get("result") or answer.get("display")
    parts = [word, answer.get("remarks")]
    return " - ".join(str(p).strip() for p in parts if str(p or "").strip())


def _join(*parts):
    return "\n".join(str(p) for p in parts if p)
